import os
import subprocess
import tempfile
import fitz  # PyMuPDF
import pytesseract
from PIL import Image
from docx import Document
import io

import json
import requests

OLLAMA_URL = "http://10.210.0.169:11434/api/generate"
# OLLAMA_URL = "http://10.210.8.133:11434/api/generate"
# OLLAMA_URL = "http://10.210.0.44:11434/api/generate"
# OLLAMA_URL = "http://10.210.8.63:11434/api/generate"
OLLAMA_MODEL = "qwen2.5:7b-instruct"
# OLLAMA_MODEL = "qwen3:8b"
# OLLAMA_MODEL = "mistral:7b-instruct"
# OLLAMA_MODEL = "gemma2:9b"

MIN_TEXT_LEN_THRESHOLD = 20

TESSERACT_LANG_MAP = {
    'en': 'eng',
    'hi': 'hin',
    'mr': 'mar',
}

def extract_text_from_image(image: Image.Image, medium='en') -> str:
    base_lang = TESSERACT_LANG_MAP.get(medium, 'eng')
    lang_str = base_lang if base_lang == 'eng' else f'{base_lang}+eng'
    print("languge: ", lang_str)
    return pytesseract.image_to_string(image, lang=lang_str, config='--psm 6')


def extract_text_from_pdf(file_path: str, medium='en') -> str:
    doc = fitz.open(file_path)
    full_text = []

    for page_num in range(len(doc)):
        page = doc[page_num]
        text = page.get_text().strip()

        if len(text) < MIN_TEXT_LEN_THRESHOLD:
            pix = page.get_pixmap(dpi=300)
            img = Image.open(io.BytesIO(pix.tobytes("png")))
            text = extract_text_from_image(img, medium=medium)

        full_text.append(text)

    doc.close()
    return "\n\n".join(full_text)


def extract_text_from_docx(file_path: str, medium='en') -> str:
    doc = Document(file_path)
    parts = []

    for para in doc.paragraphs:
        if para.text.strip():
            parts.append(para.text)

    for table in doc.tables:
        for row in table.rows:
            row_text = "\t".join(cell.text.strip() for cell in row.cells)
            if row_text.strip():
                parts.append(row_text)

    text = "\n".join(parts)

    if len(text.strip()) < MIN_TEXT_LEN_THRESHOLD:
        ocr_parts = []
        for rel in doc.part.rels.values():
            if "image" in rel.reltype:
                image_bytes = rel.target_part.blob
                img = Image.open(io.BytesIO(image_bytes))
                ocr_parts.append(extract_text_from_image(img, medium=medium))
        text = "\n\n".join(ocr_parts) if ocr_parts else text

    return text


def convert_doc_to_docx(file_path: str) -> str:
    """Convert legacy .doc to .docx using LibreOffice headless mode."""
    out_dir = tempfile.mkdtemp()
    result = subprocess.run(
        [
            "soffice", "--headless", "--convert-to", "docx",
            "--outdir", out_dir, file_path
        ],
        capture_output=True, text=True, timeout=60
    )
    if result.returncode != 0:
        raise RuntimeError(f"LibreOffice conversion failed: {result.stderr}")

    base_name = os.path.splitext(os.path.basename(file_path))[0]
    converted_path = os.path.join(out_dir, f"{base_name}.docx")

    if not os.path.exists(converted_path):
        raise RuntimeError("Converted .docx file not found")

    return converted_path


def extract_text_from_doc(file_path: str) -> str:
    converted_path = convert_doc_to_docx(file_path)
    try:
        return extract_text_from_docx(converted_path)
    finally:
        # cleanup temp converted file
        try:
            os.remove(converted_path)
            os.rmdir(os.path.dirname(converted_path))
        except OSError:
            pass


def extract_text_from_file(file_path: str, content_type: str = None, medium='en') -> str:
    ext = os.path.splitext(file_path)[1].lower()

    if ext == '.pdf':
        return extract_text_from_pdf(file_path, medium=medium)
    elif ext in ('.jpg', '.jpeg', '.png'):
        img = Image.open(file_path)
        return extract_text_from_image(img, medium=medium)
    elif ext == '.docx':
        return extract_text_from_docx(file_path, medium=medium)  # Unicode text layer, medium irrelevant
    elif ext == '.doc':
        return extract_text_from_doc(file_path)
    else:
        raise ValueError(f"Unsupported file type: {ext}")
    




PARSE_SYSTEM_PROMPT = """You convert raw exam-question text (often OCR output, possibly noisy, \
in English/Hindi/Marathi) into structured JSON for a question bank system.

Valid question_type values: SA, LA, MCQ, TF, FIB, MTF, MSQ
- SA = Short Answer, LA = Long Answer -> no options
- MCQ = single correct option, TF = True/False (exactly 2 options: True/False)
- MSQ = multiple correct options allowed
- FIB = Fill in the blank -> put "______" in question_text wherever a blank belongs
- MTF = Match the following -> use match_pairs, not options

Do NOT attempt to determine which option is correct. Only extract the question text and \
the list of option texts as written in the source.

Return ONLY this JSON shape, no markdown fences, no commentary, no explanation:
{
  "questions": [
    {
      "question_type": "MCQ",
      "question_text": "string",
      "marks": 1,
      "options": [{"text": "string"}],
      "match_pairs": [{"a": "string", "b": "string"}]
    }
  ]
}
Omit "options" for SA/LA/FIB/MTF. Omit "match_pairs" for anything except MTF.
"""

def parse_questions_with_ai(text: str, medium: str = 'en') -> list[dict]:
    prompt = f"{PARSE_SYSTEM_PROMPT}\n\nMedium/language: {medium}\n\nRaw text:\n{text}"

    response = requests.post(
        OLLAMA_URL,
        json={
            "model": OLLAMA_MODEL,
            "prompt": prompt,
            "stream": False,
            "format": "json",       # forces valid JSON output — important for small models
            "options": {"temperature": 0.1},
        },
        timeout=500,  # CPU inference can be slow, give it room
    )
    response.raise_for_status()
    raw = response.json()["response"].strip()

    parsed = json.loads(raw)  # let caller handle JSONDecodeError
    return parsed.get("questions", [])


# def parse_questions_with_ai(text: str, medium: str = 'en') -> list[dict]:
#     schema = {
#         "type": "object",
#         "properties": {
#             "questions": {
#                 "type": "array",
#                 "items": {
#                     "type": "object",
#                     "properties": {
#                         "question_type": {
#                             "type": "string",
#                             "enum": ["SA", "LA", "MCQ", "TF", "FIB", "MTF", "MSQ"],
#                         },
#                         "question_text": {"type": "string"},
#                         "marks": {"type": ["integer", "null"]},
#                         "options": {
#                             "type": "array",
#                             "items": {
#                                 "type": "object",
#                                 "properties": {"text": {"type": "string"}},
#                                 "required": ["text"],
#                             },
#                         },
#                         "match_pairs": {
#                             "type": "array",
#                             "items": {
#                                 "type": "object",
#                                 "properties": {
#                                     "a": {"type": "string"},
#                                     "b": {"type": "string"},
#                                 },
#                                 "required": ["a", "b"],
#                             },
#                         },
#                     },
#                     "required": ["question_type", "question_text", "marks"],
#                 },
#             }
#         },
#         "required": ["questions"],
#     }

#     response = requests.post(
#         OLLAMA_URL,  # e.g. f"{OLLAMA_HOST}/api/chat"
#         json={
#             "model": OLLAMA_MODEL,
#             "messages": [
#                 {"role": "system", "content": PARSE_SYSTEM_PROMPT + "\n/no_think"},
#                 {"role": "user", "content": f"Medium/language: {medium}\n\nRaw text:\n{text}"},
#             ],
#             "stream": False,
#             "format": schema,       # schema-constrained decoding, not just "json"
#             "options": {"temperature": 0.1},
#         },
#         timeout=500,
#     )
#     response.raise_for_status()
#     raw = response.json()["message"]["content"].strip()

#     parsed = json.loads(raw)  # let caller handle JSONDecodeError
#     return parsed.get("questions", [])